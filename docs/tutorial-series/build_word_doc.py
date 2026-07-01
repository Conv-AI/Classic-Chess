"""Build a single formatted Word document from the tutorial-series markdown files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = SCRIPT_DIR / "Convai-Classic-Chess-Tutorial-Series.docx"

FILES_IN_ORDER = [
    SCRIPT_DIR / "chapter-1-convai-powered-chess-coach.md",
    SCRIPT_DIR / "chapter-2-creating-the-convai-coach.md",
    SCRIPT_DIR / "chapter-3-live-interaction-and-avatar-presence.md",
    SCRIPT_DIR / "chapter-4-personalization-memory-and-custom-coaches.md",
]

SPEAKING_NOTES_HEADER = re.compile(
    r"^\*\*(?:WHAT TO SAY|Speaking notes)\*\*:?\s*$",
    re.IGNORECASE,
)


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_formatted_runs(paragraph, text: str, base_font_name: str = "Calibri", base_size: int = 11) -> None:
    pos = 0
    pattern = re.compile(
        r"(\[([^\]]+)\]\(([^)]+)\)|`([^`]+)`|\*\*([^*]+)\*\*)"
    )
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.name = base_font_name
            run.font.size = Pt(base_size)

        if match.group(2) and match.group(3):
            add_hyperlink(paragraph, match.group(2), match.group(3))
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
        elif match.group(5):
            run = paragraph.add_run(match.group(5))
            run.bold = True
            run.font.name = base_font_name
            run.font.size = Pt(base_size)

        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = base_font_name
        run.font.size = Pt(base_size)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level in range(1, 4):
        style_name = f"Heading {level}"
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        if level == 1:
            style.font.size = Pt(22)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            style.font.size = Pt(16)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(8)
        else:
            style.font.size = Pt(13)
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = parse_table_row(stripped)
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells if cell)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_formatted_runs(paragraph, cell_text, base_size=10)
            for run in paragraph.runs:
                run.font.size = Pt(10)
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")
                for run in paragraph.runs:
                    run.bold = True


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    set_paragraph_shading(paragraph, "F2F2F2")
    text = "\n".join(lines)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def normalize_speaking_notes_lines(lines: list[str]) -> list[str]:
    normalized: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped or SPEAKING_NOTES_HEADER.match(stripped):
            continue
        normalized.append(stripped)
    return normalized


def add_speaking_notes_block(doc: Document, lines: list[str]) -> None:
    body_lines = normalize_speaking_notes_lines(lines)
    if not body_lines:
        return

    container = doc.add_paragraph()
    container.paragraph_format.space_before = Pt(8)
    container.paragraph_format.space_after = Pt(8)
    container.paragraph_format.left_indent = Inches(0.2)
    container.paragraph_format.right_indent = Inches(0.2)
    set_paragraph_shading(container, "FFF2CC")

    label = container.add_run("Speaking notes\n")
    label.bold = True
    label.font.name = "Calibri"
    label.font.size = Pt(11)
    label.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)

    body = "\n".join(body_lines).strip()
    if body.startswith('"') and body.endswith('"'):
        body = body[1:-1]
    run = container.add_run(body)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    add_formatted_runs(paragraph, text)


def add_numbered(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    add_formatted_runs(paragraph, text)


def process_markdown(doc: Document, content: str) -> None:
    lines = content.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    in_blockquote = False
    blockquote_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if in_blockquote:
            if stripped.startswith(">"):
                quote_line = stripped[1:].strip()
                if quote_line:
                    blockquote_lines.append(quote_line)
                i += 1
                continue
            add_speaking_notes_block(doc, blockquote_lines)
            blockquote_lines = []
            in_blockquote = False
            continue

        if table_rows:
            if stripped.startswith("|"):
                if is_table_separator(stripped):
                    i += 1
                    continue
                table_rows.append(parse_table_row(stripped))
                i += 1
                continue
            add_table(doc, table_rows)
            table_rows = []
            continue

        if stripped.startswith("```"):
            in_code = True
            i += 1
            continue

        if stripped.startswith(">"):
            in_blockquote = True
            quote_line = stripped[1:].strip()
            if quote_line:
                blockquote_lines.append(quote_line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if not is_table_separator(stripped):
                table_rows.append(parse_table_row(stripped))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.*)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1).replace("\t", "    "))
            level = indent // 2
            add_bullet(doc, bullet_match.group(2), level=level)
            i += 1
            continue

        numbered_match = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if numbered_match:
            indent = len(numbered_match.group(1).replace("\t", "    "))
            level = indent // 2
            add_numbered(doc, numbered_match.group(2), level=level)
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_formatted_runs(paragraph, stripped)
        i += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)
    if in_blockquote and blockquote_lines:
        add_speaking_notes_block(doc, blockquote_lines)
    if table_rows:
        add_table(doc, table_rows)


def build_document() -> Path:
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    configure_styles(doc)

    for idx, file_path in enumerate(FILES_IN_ORDER):
        if idx > 0:
            doc.add_page_break()
        content = file_path.read_text(encoding="utf-8")
        process_markdown(doc, content)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(f"Created: {output}")
